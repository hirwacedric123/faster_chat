import os
import re
from typing import List, Tuple
from django.conf import settings
import pypdf
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

from .models import Document, DocumentChunk

class DocumentProcessor:
    """Utility class for processing uploaded documents"""
    
    def __init__(self, document: Document):
        self.document = document
        self.file_path = os.path.join(settings.MEDIA_ROOT, self.document.file.name)
        self.chunks = []
    
    def process(self) -> bool:
        """Process the document and split into chunks"""
        try:
            # Extract text based on file type
            text = self.extract_text()
            
            # Split text into chunks
            chunks = self.split_into_chunks(text)
            
            # Save chunks to database
            for i, chunk_text in enumerate(chunks):
                DocumentChunk.objects.create(
                    document=self.document,
                    content=chunk_text,
                    chunk_number=i
                )
            
            return True
        except Exception as e:
            self.document.error_message = str(e)
            self.document.status = 'failed'
            self.document.save()
            return False
    
    def extract_text(self) -> str:
        """Extract text from the document based on file type"""
        extension = self.document.file_extension
        
        if extension == '.pdf':
            return self._extract_from_pdf()
        elif extension in ['.doc', '.docx']:
            return self._extract_from_docx()
        elif extension in ['.txt', '.md']:
            return self._extract_from_text()
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _extract_from_pdf(self) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_reader = pypdf.PdfReader(self.file_path)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n\n"
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        
        return text
    
    def _extract_from_docx(self) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(self.file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " | "
                    text += row_text.strip(" | ") + "\n"
                text += "\n"
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
        
        return text
    
    def _extract_from_text(self) -> str:
        """Extract text from plain text file"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings if utf-8 fails
            encodings = ['latin-1', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(self.file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with supported encodings")
        except Exception as e:
            raise Exception(f"Error extracting text from text file: {str(e)}")
    
    def split_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks for processing"""
        # Clean the text
        text = self._clean_text(text)
        
        # Use LangChain's text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,          # Increased from 1000 to 1500 for better context handling
            chunk_overlap=150,        # Increased from 100 to 150 for better overlap
            length_function=len,      # Function to measure length
            separators=["\n\n", "\n", ". ", " ", ""]  # Added period+space as a separator
        )
        
        return text_splitter.split_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        
        # Strip whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text 